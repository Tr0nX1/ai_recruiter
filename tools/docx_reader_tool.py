"""DOCX text extraction."""

from __future__ import annotations

import os

from crewai.tools import BaseTool
from docx import Document
from pydantic import BaseModel, Field


class DOCXReaderInput(BaseModel):
    file_path: str = Field(description="Path to a .docx file")


class DOCXReaderTool(BaseTool):
    name: str = "DOCXReaderTool"
    description: str = "Reads a DOCX Word document and extracts paragraph and table text."
    args_schema: type[BaseModel] = DOCXReaderInput

    def _run(self, file_path: str) -> dict:
        if not os.path.isfile(file_path):
            return {
                "raw_text": "",
                "page_count": 0,
                "extraction_status": "failed",
                "error_message": f"File not found: {file_path}",
            }
        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables: list[str] = []
            for table in doc.tables:
                for row in table.rows:
                    tables.append(" | ".join(cell.text for cell in row.cells))
            raw_text = "\n".join(paragraphs + tables)
            stripped = raw_text.strip()
            return {
                "raw_text": stripped,
                "page_count": 1,
                "extraction_status": "success" if stripped else "failed",
                "error_message": None,
            }
        except Exception as e:  # noqa: BLE001
            return {
                "raw_text": "",
                "page_count": 0,
                "extraction_status": "failed",
                "error_message": str(e),
            }
